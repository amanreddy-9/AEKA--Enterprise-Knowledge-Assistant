"""Document Processor — PDF & DOCX extraction + chunking"""

from pathlib import Path
from typing import List, Dict
from langchain_text_splitters import RecursiveCharacterTextSplitter


class DocumentProcessor:
    def __init__(self, chunk_size: int = 512, chunk_overlap: int = 64):
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", ". ", " ", ""],
        )

    def process_file(self, file_path: str, original_name: str) -> List[Dict]:
        ext = Path(file_path).suffix.lower()
        try:
            if ext == ".pdf":
                raw_text = self._extract_pdf(file_path)
            elif ext == ".docx":
                raw_text = self._extract_docx(file_path)
            else:
                return []

            if not raw_text.strip():
                return []

            return self._chunk(raw_text, original_name)
        except Exception as e:
            print(f"[Processor] Error processing {original_name}: {e}")
            return []

    def _extract_pdf(self, path: str) -> str:
        import fitz
        import os

        print(f"[PDF] Opening: {path}")
        print(f"[PDF] File exists: {os.path.exists(path)}")

        doc = fitz.open(path)
        print(f"[PDF] Pages: {len(doc)}")

        pages = []
        for page in doc:
            text = page.get_text("text")
            print(f"[PDF] Page chars: {len(text)}")
            pages.append(text)

        doc.close()

        result = "\n\n".join(pages)
        print(f"[PDF] Total extracted chars: {len(result)}")
        return result

    def _extract_docx(self, path: str) -> str:
        from docx import Document

        doc = Document(path)

        parts = [p.text for p in doc.paragraphs if p.text.strip()]

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())

        return "\n\n".join(parts)

    def _chunk(self, text: str, source: str) -> List[Dict]:
        chunks = self.splitter.split_text(text)
        return [
            {"text": c, "file": source, "chunk_id": i}
            for i, c in enumerate(chunks)
            if c.strip()
        ]